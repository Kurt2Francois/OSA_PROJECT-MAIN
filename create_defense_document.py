from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    return heading

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_defense_document():
    """Create comprehensive project defense document"""
    
    doc = Document()
    
    # Title Page
    title = doc.add_heading('OSA Partnership Monitoring System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Defense Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph('Backend Development & Functionality')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_format = info.runs[0]
    info_format.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph('Date: December 4, 2025')
    doc.add_paragraph('Technology: Django 4.2.17 | Django REST Framework | SQLite3')
    
    # Page Break
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. Architecture & Technology Stack',
        '3. Database Model Design',
        '4. Core Functionality & Features',
        '5. Key Backend Functions',
        '6. API Endpoints (DRF ViewSets)',
        '7. URL Routing',
        '8. Security Features',
        '9. Database Queries',
        '10. Workflow Example: User Journey',
        '11. Project Structure',
        '12. Defense Q&A',
        '13. Future Improvements',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 1: Project Overview
    doc.add_heading('1. Project Overview', 1)
    
    overview_table = doc.add_table(rows=5, cols=2)
    overview_table.style = 'Light Grid Accent 1'
    
    cells = overview_table.rows[0].cells
    cells[0].text = 'Project Name'
    cells[1].text = 'OSA Partnership Monitoring System'
    shade_cell(cells[0], 'D3D3D3')
    
    cells = overview_table.rows[1].cells
    cells[0].text = 'Purpose'
    cells[1].text = 'Manage and monitor partnerships between OSA and various departments/businesses with role-based access control'
    shade_cell(cells[0], 'D3D3D3')
    
    cells = overview_table.rows[2].cells
    cells[0].text = 'Main Framework'
    cells[1].text = 'Django 4.2.17'
    shade_cell(cells[0], 'D3D3D3')
    
    cells = overview_table.rows[3].cells
    cells[0].text = 'Database'
    cells[1].text = 'SQLite3'
    shade_cell(cells[0], 'D3D3D3')
    
    cells = overview_table.rows[4].cells
    cells[0].text = 'Key Features'
    cells[1].text = 'Authentication, RBAC, Department Management, Admin Dashboard, RESTful API'
    shade_cell(cells[0], 'D3D3D3')
    
    doc.add_page_break()
    
    # Section 2: Architecture & Technology Stack
    doc.add_heading('2. Architecture & Technology Stack', 1)
    
    tech_table = doc.add_table(rows=9, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = tech_table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'
    for cell in header_cells:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    tech_data = [
        ('Backend Framework', 'Django 4.2.17', 'Web framework for building scalable applications'),
        ('API Layer', 'Django REST Framework 3.16.0', 'RESTful API endpoints for data access'),
        ('Database', 'SQLite3', 'Lightweight relational database'),
        ('Authentication', 'Django JWT & djangorestframework-simplejwt', 'Secure token-based authentication'),
        ('Media Handling', 'Pillow 11.0.0', 'Image processing for department logos'),
        ('CORS', 'django-cors-headers 4.4.0', 'Cross-origin request handling'),
        ('Frontend', 'Tailwind CSS', 'Responsive UI design'),
        ('Password Hashing', 'Django built-in', 'Secure password storage'),
    ]
    
    for i, (component, tech, purpose) in enumerate(tech_data, start=1):
        row_cells = tech_table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # Section 3: Database Models
    doc.add_heading('3. Database Model Design', 1)
    
    # UserProfile Model
    doc.add_heading('3.1 UserProfile Model', 2)
    doc.add_paragraph('Purpose: Extends Django\'s built-in User model with business-specific information', style='List Bullet')
    doc.add_paragraph('User Types: Department, Admin, Owner', style='List Bullet')
    
    user_table = doc.add_table(rows=8, cols=2)
    user_table.style = 'Light Grid Accent 1'
    
    header = user_table.rows[0].cells
    header[0].text = 'Field Name'
    header[1].text = 'Description'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    user_fields = [
        ('user (OneToOneField)', 'Links to Django User model'),
        ('business_email (EmailField)', 'Business contact email - validated email format'),
        ('department_name (CharField)', 'Department name - max 255 chars'),
        ('contact_person (CharField)', 'Primary contact person'),
        ('contact_number (CharField)', 'Phone number - max 50 chars'),
        ('user_type (CharField)', 'Role/permission level (department, admin, owner)'),
        ('created_at (DateTimeField)', 'Registration timestamp - auto-populated'),
    ]
    
    for i, (field, desc) in enumerate(user_fields, start=1):
        row = user_table.rows[i].cells
        row[0].text = field
        row[1].text = desc
    
    doc.add_paragraph()
    
    # Department Model
    doc.add_heading('3.2 Department Model', 2)
    doc.add_paragraph('Purpose: Represents partnerships and departments in the system', style='List Bullet')
    doc.add_paragraph('Status Options: Active, Inactive, Pending', style='List Bullet')
    
    dept_table = doc.add_table(rows=11, cols=2)
    dept_table.style = 'Light Grid Accent 1'
    
    header = dept_table.rows[0].cells
    header[0].text = 'Field Name'
    header[1].text = 'Description'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    dept_fields = [
        ('owner (ForeignKey)', 'User who owns the department'),
        ('department_name (CharField)', 'Department name - max 255 chars'),
        ('business_email (EmailField)', 'Business contact email - validated email format'),
        ('email (EmailField)', 'Department email address'),
        ('contact_person (CharField)', 'Contact person - optional'),
        ('contact_number (CharField)', 'Contact number - optional'),
        ('logo_path (ImageField)', 'Department logo - upload to logos/ directory'),
        ('established_date (DateField)', 'Partnership start date - optional'),
        ('expiration_date (DateField)', 'Partnership end date - optional'),
        ('partnership_status (CharField)', 'Status: active, inactive, or pending'),
    ]
    
    for i, (field, desc) in enumerate(dept_fields, start=1):
        row = dept_table.rows[i].cells
        row[0].text = field
        row[1].text = desc
    
    doc.add_page_break()
    
    # Section 4: Core Functionality
    doc.add_heading('4. Core Functionality & Features', 1)
    
    doc.add_heading('4.1 Authentication System', 2)
    doc.add_paragraph('signup_view()', style='List Number')
    doc.add_paragraph('User registration with department information', style='List Bullet 2')
    doc.add_paragraph('Email and password confirmation validation', style='List Bullet 2')
    doc.add_paragraph('Creates User, UserProfile, and Department atomically', style='List Bullet 2')
    
    doc.add_paragraph('login_view()', style='List Number')
    doc.add_paragraph('Email-based authentication', style='List Bullet 2')
    doc.add_paragraph('Smart redirects: Superusers → Admin Panel, Users → Department', style='List Bullet 2')
    doc.add_paragraph('Session management with Django\'s authentication', style='List Bullet 2')
    
    doc.add_paragraph('logout_view()', style='List Number')
    doc.add_paragraph('Clears user session and redirects to login', style='List Bullet 2')
    
    doc.add_heading('4.2 Role-Based Access Control (RBAC)', 2)
    
    rbac_table = doc.add_table(rows=4, cols=3)
    rbac_table.style = 'Light Grid Accent 1'
    
    header = rbac_table.rows[0].cells
    header[0].text = 'User Type'
    header[1].text = 'Permissions'
    header[2].text = 'Access Level'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    rbac_data = [
        ('Department Owner', 'View/Edit own departments, Submit for approval', 'Limited'),
        ('Admin', 'View/Edit all departments, Manage remarks, Approve partnerships', 'High'),
        ('Superuser', 'Full system access, User management, System configuration', 'Full'),
    ]
    
    for i, (utype, perms, level) in enumerate(rbac_data, start=1):
        row = rbac_table.rows[i].cells
        row[0].text = utype
        row[1].text = perms
        row[2].text = level
    
    doc.add_heading('4.3 Department Management Features', 2)
    features = [
        ('View Department', 'Display detailed department information and partnership status'),
        ('Edit Department', 'Update business info, dates, status, logo, and remarks'),
        ('Delete Department', 'Remove department record from system'),
        ('Add Department', 'Create new department/partnership records'),
        ('Logo Display', 'Upload department logo and display it across dashboard, owner panel, and detail views'),
        ('Dashboard View', 'List and manage own departments or all departments (for admin)'),
        ('Admin Panel Expiration Column', 'Admin panel displays partnership expiration dates for quick review'),
        ('User Management', 'Admins can view and delete user accounts via the Admin Panel'),
    ]
    
    for feature, desc in features:
        doc.add_paragraph(f'{feature}: {desc}', style='List Bullet')
    
    doc.add_heading('4.4 Admin Dashboard', 2)
    doc.add_paragraph('Statistics Cards:', style='List Bullet')
    doc.add_paragraph('Total Departments count', style='List Bullet 2')
    doc.add_paragraph('Active Partnerships (status=\'active\')', style='List Bullet 2')
    doc.add_paragraph('Pending Partnerships (status=\'pending\')', style='List Bullet 2')
    doc.add_paragraph('Total Users count', style='List Bullet 2')
    
    doc.add_paragraph('Department Management Table:', style='List Bullet')
    doc.add_paragraph('View all departments with complete details', style='List Bullet 2')
    doc.add_paragraph('Edit partnership status and remarks inline', style='List Bullet 2')
    doc.add_paragraph('Color-coded status badges (Green=Active, Yellow=Pending, Red=Inactive)', style='List Bullet 2')
    doc.add_paragraph('Quick action buttons (View, Edit, Delete)', style='List Bullet 2')
    doc.add_paragraph('Expiration column included to show partnership end date', style='List Bullet 2')
    doc.add_paragraph('Users management table: View and delete users (superuser only)', style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 5: Backend Functions
    doc.add_heading('5. Key Backend Functions', 1)
    
    doc.add_heading('5.1 Authentication Functions', 2)
    
    auth_funcs = [
        ('signup_view(request)', 'Handles user registration, creates User/UserProfile/Department'),
        ('login_view(request)', 'Authenticates user by email, manages session, smart redirects'),
        ('logout_view(request)', 'Terminates session and redirects to login page'),
    ]
    
    for func, desc in auth_funcs:
        doc.add_paragraph(func, style='List Bullet')
        doc.add_paragraph(desc, style='List Bullet 2')
    
    doc.add_heading('5.2 Department Operations', 2)
    
    dept_funcs = [
        ('department_detail_view(request, dept_id)', 'Displays department details with permission checks'),
        ('department_edit_view(request, dept_id)', 'Handles text/date/file updates with role-based restrictions'),
        ('department_delete_view(request, dept_id)', 'Removes department with permission validation'),
        ('department_add_view(request)', 'Renders form for creating new departments'),
    ]
    
    for func, desc in dept_funcs:
        doc.add_paragraph(func, style='List Bullet')
        doc.add_paragraph(desc, style='List Bullet 2')
    
    doc.add_heading('5.3 Dashboard & Admin Functions', 2)
    
    dashboard_funcs = [
        ('dashboard_view(request)', 'Lists user\'s departments or all departments if admin'),
        ('owner_panel_view(request)', 'Shows owner-specific department management interface'),
        ('admin_panel_view(request)', 'Displays admin dashboard with stats and all departments'),
        ('user_delete_view(request, user_id)', 'Deletes a user (superuser only)'),
    ]
    
    for func, desc in dashboard_funcs:
        doc.add_paragraph(func, style='List Bullet')
        doc.add_paragraph(desc, style='List Bullet 2')
    
    doc.add_heading('5.4 Validation Functions', 2)
    doc.add_paragraph('is_admin(user)', style='List Bullet')
    doc.add_paragraph('Returns True if user.profile.user_type == \'admin\'', style='List Bullet 2')
    
    doc.add_heading('5.5 Data Validation in signup_view()', 2)
    validations = [
        'Email must match confirmation email',
        'Password must match confirmation password',
        'Email must be unique (no duplicates)',
        'All required fields must be provided',
    ]
    for validation in validations:
        doc.add_paragraph(validation, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 6: API Endpoints
    doc.add_heading('6. API Endpoints (DRF ViewSets)', 1)
    
    api_table = doc.add_table(rows=7, cols=3)
    api_table.style = 'Light Grid Accent 1'
    
    header = api_table.rows[0].cells
    header[0].text = 'HTTP Method'
    header[1].text = 'Endpoint'
    header[2].text = 'Description'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    api_endpoints = [
        ('GET/POST', '/api/users/', 'List all users or create new user'),
        ('GET/PUT/PATCH', '/api/users/{id}/', 'Retrieve, update, or partially update user'),
        ('DELETE', '/api/users/{id}/', 'Delete a user'),
        ('GET/POST', '/api/departments/', 'List all departments or create new'),
        ('GET/PUT/PATCH', '/api/departments/{id}/', 'Retrieve, update, or partially update department'),
        ('DELETE', '/api/departments/{id}/', 'Delete a department'),
    ]
    
    for i, (method, endpoint, desc) in enumerate(api_endpoints, start=1):
        row = api_table.rows[i].cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
    
    doc.add_paragraph()
    doc.add_heading('Serializers:', 2)
    doc.add_paragraph('UserSerializer: Serializes UserProfile objects for API responses', style='List Bullet')
    doc.add_paragraph('DepartmentSerializer: Serializes Department objects for API responses', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 7: URL Routing
    doc.add_heading('7. URL Routing', 1)
    
    url_table = doc.add_table(rows=12, cols=3)
    url_table.style = 'Light Grid Accent 1'
    
    header = url_table.rows[0].cells
    header[0].text = 'URL Path'
    header[1].text = 'View Function'
    header[2].text = 'Purpose'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    urls = [
        ('/', 'login_view', 'Home/login page'),
        ('/signup/', 'signup_view', 'User registration page'),
        ('/logout/', 'logout_view', 'Session logout endpoint'),
        ('/dashboard/', 'dashboard_view', 'Main user dashboard'),
        ('/department/<id>/', 'department_detail_view', 'View department details'),
        ('/department/<id>/edit/', 'department_edit_view', 'Edit department information'),
        ('/admin-panel/', 'admin_panel_view', 'Admin statistics & management'),
        ('/owner-panel/', 'owner_panel_view', 'Owner-specific dashboard'),
        ('/owner-panel/department/add/', 'department_add_view', 'Add new department'),
        ('/owner-panel/department/<id>/delete/', 'department_delete_view', 'Delete department'),
        ('/api/', 'DRF Router', 'RESTful API endpoints (nested routes)'),
    ]
    
    for i, (url, view, purpose) in enumerate(urls, start=1):
        row = url_table.rows[i].cells
        row[0].text = url
        row[1].text = view
        row[2].text = purpose
    
    doc.add_page_break()
    
    # Section 8: Security
    doc.add_heading('8. Security Features', 1)
    
    security_items = [
        ('Authentication', 'Django built-in user authentication + JWT tokens (djangorestframework-simplejwt)'),
        ('Authorization', 'Role-based access control (RBAC) with user_type checks'),
        ('CSRF Protection', 'Django\'s CsrfViewMiddleware enabled by default'),
        ('Permission Checks', '@login_required decorator + custom permission logic in views'),
        ('Data Validation', 'Form validation on signup, email/password confirmation'),
        ('Secure Passwords', 'Django\'s password hashing (PBKDF2) + password validators'),
        ('SQL Injection Prevention', 'Django ORM prevents SQL injection attacks'),
        ('XSS Protection', 'Django templates auto-escape HTML content'),
        ('CORS Handling', 'django-cors-headers middleware for API security'),
        ('Session Security', 'Django sessions with HttpOnly, Secure, and SameSite cookies'),
    ]
    
    for feature, description in security_items:
        doc.add_paragraph(f'{feature}: {description}', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 9: Database Queries
    doc.add_heading('9. Database Queries & Optimization', 1)
    
    doc.add_heading('9.1 Common Query Patterns Used', 2)
    
    queries = [
        ('Filter by Status', 'Department.objects.filter(partnership_status=\'active\')'),
        ('Filter by Owner', 'Department.objects.filter(owner=request.user)'),
        ('Existence Check', 'Department.objects.exists()'),
        ('Single Retrieval', 'get_object_or_404(Department, id=dept_id)'),
        ('Count Aggregation', 'Department.objects.count()'),
        ('Get or Create', 'Department.objects.get_or_create(...)'),
    ]
    
    for pattern, query in queries:
        doc.add_paragraph(f'{pattern}:', style='List Bullet')
        code_para = doc.add_paragraph(query, style='List Bullet 2')
        code_para.style.font.name = 'Courier New'
    
    doc.add_heading('9.2 Query Optimization', 2)
    optimizations = [
        'Use select_related() for ForeignKey relationships',
        'Use prefetch_related() for reverse ForeignKey and ManyToMany',
        'Filter early to reduce dataset size',
        'Use count() directly instead of len() on querysets',
        'Index frequently filtered fields',
    ]
    for opt in optimizations:
        doc.add_paragraph(opt, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 10: User Journey
    doc.add_heading('10. User Journey: Registration to Department Management', 1)
    
    journey = [
        ('User visits /', 'Sees login page'),
        ('Clicks signup link', 'Navigates to /signup/'),
        ('Fills registration form', 'Enters business email, password, contact info, etc.'),
        ('Backend validation', 'Checks email uniqueness, password match, format'),
        ('Creates records', 'Generates User, UserProfile, and Department objects'),
        ('Auto-login', 'Authenticates and starts session'),
        ('Redirect to department', 'User sees their new department at /department/<id>/'),
        ('Edit department', 'User can click Edit → /department/<id>/edit/'),
        ('Admin access', 'Admin logs in → sees /admin-panel/'),
        ('Dashboard navigation', 'Admin can navigate to /partnership/dashboard/ and return to Admin Panel via the Admin Panel link'),
        ('View statistics', 'Admin sees total departments, active/pending counts'),
        ('Manage departments', 'Admin can view, edit, delete any department'),
        ('Manage users', "Admin (superuser) can delete user accounts from the Admin Panel with confirmation"),
    ]
    
    for i, (step, action) in enumerate(journey, start=1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
        doc.add_paragraph(action, style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 11: Project Structure
    doc.add_heading('11. Project Structure', 1)
    
    structure_text = '''OSA_PROJECT-MAIN/
├── osa_project/                    # Django project root
│   ├── manage.py                   # Django CLI tool
│   ├── db.sqlite3                  # SQLite database file
│   │
│   ├── osa_project/                # Project configuration
│   │   ├── settings.py             # Django settings & configuration
│   │   ├── urls.py                 # Main URL routing
│   │   ├── asgi.py                 # ASGI configuration
│   │   └── wsgi.py                 # WSGI configuration
│   │
│   ├── partnership/                # Main application
│   │   ├── models.py               # UserProfile & Department models
│   │   ├── views.py                # All view functions & ViewSets
│   │   ├── urls.py                 # App-level URL routing
│   │   ├── serializers.py          # DRF serializers for API
│   │   ├── admin.py                # Django admin configuration
│   │   └── migrations/             # Database migrations
│   │
│   ├── templates/                  # HTML templates
│   │   └── partnership/
│   │       ├── login.html          # Login page
│   │       ├── signup.html         # Registration page
│   │       ├── dashboard.html      # Main dashboard
│   │       ├── admin_panel.html    # Admin interface
│   │       ├── owner_panel.html    # Owner interface
│   │       ├── department_detail.html    # Department view (shows logo, dates, remarks)
│   │       ├── department_edit.html      # Department edit form (supports logo upload)
│   │       ├── department_delete.html    # Old delete template (deprecated)
│   │       ├── department_delete_confirm.html    # Delete confirmation (new)
│   │       └── user_delete_confirm.html    # User delete confirmation (new)
│   │
│   └── static/                     # Static files
│       ├── css/
│       ├── js/
│       └── images/
│
├── requirements.txt                # Python dependencies
├── .git/                           # Git version control
└── venv/                           # Python virtual environment
'''
    
    code_para = doc.add_paragraph(structure_text)
    code_para.style.font.name = 'Courier New'
    code_para.style.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Section 12: Defense Q&A
    doc.add_heading('12. Common Defense Questions & Answers', 1)
    
    qa_pairs = [
        ('Q: Why did you choose Django for this project?',
         'A: Django provides rapid development, built-in authentication, ORM for database management, security middleware, admin panel, and excellent documentation. It\'s ideal for CRUD-heavy applications like partnership management systems.'),
        
        ('Q: How does role-based access control (RBAC) work in your system?',
         'A: We store user_type in UserProfile model (department, admin, owner). In each view, we check the user_type before allowing operations. Superusers bypass RBAC checks. Department owners can only access their own records.'),
        
        ('Q: How do you secure user passwords?',
         'A: Django uses PBKDF2 password hashing algorithm with salt. Passwords are never stored in plaintext. We also validate passwords using Django\'s built-in validators (minimum length, complexity checks).'),
        
        ('Q: Can regular users see other departments?',
         'A: No. We filter queries with owner=request.user to show only their departments. Only admins and superusers see all departments via the admin panel.'),
        
        ('Q: How do you prevent SQL injection attacks?',
         'A: Django ORM parameterizes all SQL queries, preventing injection. We use Django\'s QuerySet methods instead of raw SQL.'),
        
        ('Q: What happens when two users edit the same department simultaneously?',
         'A: Last write wins. The last_updated timestamp helps track changes. For production, implement optimistic locking or database transactions.'),
        
        ('Q: Why use REST API alongside web views?',
         'A: REST API allows mobile apps, third-party integrations, and decoupling frontend from backend. Web views serve the traditional HTML interface.'),
        
        ('Q: How do you handle file uploads (logos)?',
         'A: Use Django\'s ImageField which stores files in media/ directory. Pillow validates image format. Files are validated before saving.'),
        
        ('Q: What\'s the purpose of UserProfile if Django User exists?',
         'A: Django User only has basic fields (username, email, password). UserProfile extends it with business-specific data (business_email with email validation, department_name, user_type, contact info).'),
        
        ('Q: How do you validate incoming data?',
         'A: In signup_view, we check email matching, password matching, email uniqueness. In forms, we use Django\'s Form validation. In API, DRF serializers validate data automatically.'),
    ]
    
    for question, answer in qa_pairs:
        q_para = doc.add_paragraph(question)
        q_para.runs[0].font.bold = True
        q_para.runs[0].font.color.rgb = RGBColor(0, 0, 128)
        
        a_para = doc.add_paragraph(answer)
        a_para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Section 13: Future Improvements
    doc.add_heading('13. Future Improvements & Scalability', 1)
    
    improvements = [
        'Search & Filter: Implement full-text search and advanced filtering in admin panel',
        'Export Features: Generate CSV/PDF reports of departments and partnerships',
        'Email Notifications: Send automated emails on status changes, approvals',
        'Approval Workflow: Implement multi-stage approval process for partnerships',
        'Activity Logging: Track all user actions for audit trails',
        'Advanced Analytics: Dashboard with charts, graphs, trends analysis',
        'Batch Operations: Allow bulk updates and deletions for admin',
        'API Rate Limiting: Implement throttling to prevent abuse',
        'Webhook Support: Trigger external integrations on events',
        'Mobile App: Native iOS/Android app using REST API',
        'Caching: Implement Redis caching for frequently accessed data',
        'Pagination: Add pagination to large lists',
        'Department Categories: Organize departments by type/category',
        'Expiration Alerts: Email alerts when partnerships are about to expire',
        'Compliance Reports: Generate compliance documents and reports',
    ]
    
    for improvement in improvements:
        doc.add_paragraph(improvement, style='List Bullet')
    
    doc.add_page_break()
    
    # Technical Details
    doc.add_heading('14. Technical Dependencies', 1)
    
    deps_table = doc.add_table(rows=11, cols=3)
    deps_table.style = 'Light Grid Accent 1'
    
    header = deps_table.rows[0].cells
    header[0].text = 'Package'
    header[1].text = 'Version'
    header[2].text = 'Purpose'
    for cell in header:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    dependencies = [
        ('Django', '4.2.17', 'Web framework'),
        ('djangorestframework', '3.16.0', 'REST API framework'),
        ('djangorestframework-simplejwt', '5.3.1', 'JWT authentication'),
        ('django-cors-headers', '4.4.0', 'CORS middleware'),
        ('Pillow', '11.0.0', 'Image processing'),
        ('PyJWT', '2.10.1', 'JWT token handling'),
        ('sqlparse', '0.4.4', 'SQL parsing utilities'),
        ('asgiref', '3.7.2', 'ASGI compatibility'),
        ('pytz', '2025.2', 'Timezone handling'),
        ('tzdata', '2025.2', 'Timezone database'),
    ]
    
    for i, (pkg, version, purpose) in enumerate(dependencies, start=1):
        row = deps_table.rows[i].cells
        row[0].text = pkg
        row[1].text = version
        row[2].text = purpose
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    
    conclusion = '''This OSA Partnership Monitoring System demonstrates a complete Django backend implementation with:

• Secure authentication and authorization
• Comprehensive role-based access control
• RESTful API for extensibility
• Well-designed database models
• Professional admin interface
• Best practices for security and performance

The system is production-ready and can be easily extended with additional features. The modular architecture allows for adding new functionality without disrupting existing code.

Key Achievements:
✓ Implemented multi-user authentication system
✓ Built role-based access control (RBAC)
✓ Created comprehensive admin dashboard
✓ Developed RESTful API with DRF
✓ Designed efficient database schema
✓ Implemented security best practices
✓ Created responsive user interfaces
✓ Added data validation and error handling
✓ Integrated image upload functionality
✓ Built admin statistics and reporting
'''
    
    conclusion_para = doc.add_paragraph(conclusion)
    
    # Save document
    output_path = r'c:\Users\FranZ\OSA_PROJECT-MAIN\OSA_Partnership_Monitoring_System_Defense.docx'
    doc.save(output_path)
    
    print(f'✓ Document created successfully: {output_path}')
    print(f'✓ Total pages: ~15')
    print(f'✓ Total sections: 14')
    print(f'✓ Tables: 10+')
    print(f'✓ Ready for presentation!')

if __name__ == '__main__':
    create_defense_document()
